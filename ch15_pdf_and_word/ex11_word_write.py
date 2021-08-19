import docx


doc = docx.Document()
doc.add_paragraph('Hello, world!')
doc.save('helloworld.docx')

doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('multipleParagraphs.docx')

# Add headings
doc = docx.Document()
doc.add_heading('Header 0', 0)              # docx.text.Paragraph object at 0x00000000036CB3C8>
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# Add Line and Page Breaks
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

doc = docx.Document()
doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('pic.docx')